import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime
import plotly.graph_objects as go
import plotly.express as px
import requests
from io import BytesIO
import base64
import os
import io
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')
import seaborn as sns

# PDF generation (ReportLab)
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm, mm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader

# =========================
# Page Config
# =========================
st.set_page_config(
    page_title="Sample Track Analytics Dashboard",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# =========================
# Simple Clean CSS (NO emojis)
# =========================
st.markdown("""
<style>
    .app-header {
        font-size: 1.8rem;
        font-weight: 700;
        color: #0f172a;
        margin: 0.25rem 0 1.0rem 0;
        padding: 0.75rem 1rem;
        background: #f8fafc;
        border: 1px solid #e2e8f0;
        border-radius: 12px;
    }
    .section-title{
        font-size: 1.2rem;
        font-weight: 700;
        color: #0f172a;
        margin-top: 1.2rem;
        margin-bottom: 0.6rem;
    }
    .kpi-card{
        background: #ffffff;
        border: 1px solid #e2e8f0;
        border-radius: 14px;
        padding: 14px 16px;
        box-shadow: 0 1px 2px rgba(15,23,42,0.04);
        height: 118px;
    }
    .kpi-label{
        font-size: 0.85rem;
        color: #475569;
        font-weight: 600;
        margin-bottom: 6px;
    }
    .kpi-value{
        font-size: 1.8rem;
        font-weight: 800;
        color: #0f172a;
        line-height: 1.1;
        margin-bottom: 4px;
    }
    .kpi-sub{
        font-size: 0.85rem;
        color: #64748b;
        font-weight: 600;
    }
    .hint{
        background: #f8fafc;
        border: 1px solid #e2e8f0;
        border-left: 4px solid #3b82f6;
        border-radius: 10px;
        padding: 10px 12px;
        color: #334155;
        font-size: 0.9rem;
        margin-top: 8px;
        margin-bottom: 10px;
    }
    .export-section {
        background: linear-gradient(135deg, #f8fafc 0%, #f1f5f9 100%);
        border: 1px solid #e2e8f0;
        border-radius: 12px;
        padding: 1.5rem;
        margin: 1rem 0;
    }
    .export-card {
        background: white;
        border: 1px solid #e2e8f0;
        border-radius: 10px;
        padding: 1rem;
        margin: 0.5rem 0;
        transition: all 0.3s ease;
    }
    .export-card:hover {
        box-shadow: 0 4px 12px rgba(15, 23, 42, 0.1);
        border-color: #3b82f6;
    }
    .export-title {
        font-size: 1.1rem;
        font-weight: 600;
        color: #0f172a;
        margin-bottom: 0.5rem;
    }
    .export-desc {
        font-size: 0.9rem;
        color: #64748b;
        margin-bottom: 1rem;
    }
    .stButton > button {
        width: 100%;
        border-radius: 8px;
        padding: 0.75rem 1rem;
        font-weight: 600;
        transition: all 0.3s ease;
    }
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(59, 130, 246, 0.2);
    }
</style>
""", unsafe_allow_html=True)

st.markdown('<div class="app-header">Sample Track Analytics Dashboard</div>', unsafe_allow_html=True)

# =========================
# Google Sheet Config
# =========================
SPREADSHEET_KEY = "1lkztBZ4eG1BQx-52XgnA6w8YIiw-Sm85pTlQQziurfw"
WORKSHEET_NAME = "Test"

# =========================
# Text Normalization
# =========================
def norm_text(x: str) -> str:
    s = str(x).strip().lower()
    s = s.replace("_", " ").replace("'", "").replace("`", "")
    s = " ".join(s.split())

    s = s.replace("sar e pul", "sar-e-pul")
    s = s.replace("sare pul", "sar-e-pul")
    s = s.replace("maidan wardak", "wardak")
    s = s.replace("maydan wardak", "wardak")
    s = s.replace("panjshir", "panjsher")
    s = s.replace("jawzjan", "jowzjan")
    return s

def remove_total_rows(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    for c in ["Region", "Province", "District"]:
        if c in df.columns:
            df[c] = df[c].astype(str).str.strip()
            df = df[~df[c].str.contains(r"\btotal\b", case=False, na=False)]
    return df

def safe_num(series: pd.Series) -> pd.Series:
    return pd.to_numeric(series, errors="coerce").fillna(0)

def find_prefixed_metric(df_cols, prefix, keywords):
    cols = [str(c).strip() for c in df_cols]
    candidates = [c for c in cols if c.startswith(prefix)]
    for kw in keywords:
        for c in candidates:
            if kw.lower() in c.lower():
                return c
    if len(candidates) == 1:
        return candidates[0]
    return None

def build_tool_view(df_raw: pd.DataFrame, tool: str) -> pd.DataFrame:
    df = df_raw.copy()
    if df.shape[1] < 3:
        raise ValueError("Sheet must have at least 3 columns (Region, Province, District).")

    cols = list(df.columns)
    df = df.rename(columns={cols[0]: "Region", cols[1]: "Province", cols[2]: "District"})

    df["Region"] = df["Region"].astype(str).str.strip()
    df["Province"] = df["Province"].astype(str).str.strip()
    df["District"] = df["District"].astype(str).str.strip()

    df = remove_total_rows(df)

    prefix_map = {"CBE": "CBE-", "PBs": "PBs-", "Total": "Total-"}
    prefix = prefix_map[tool]

    col_target = find_prefixed_metric(df.columns, prefix, ["target", "sample", "samplesize"])
    col_received = find_prefixed_metric(df.columns, prefix, ["received", "collected"])
    col_approved = find_prefixed_metric(df.columns, prefix, ["approved", "approve"])
    col_pending = find_prefixed_metric(df.columns, prefix, ["pending", "review"])
    col_rejected = find_prefixed_metric(df.columns, prefix, ["rejected", "reject"])
    col_checked = find_prefixed_metric(df.columns, prefix, ["checked", "reviewed"])
    
    # Add columns for Unable to visit
    col_unable_visit = find_prefixed_metric(df.columns, prefix, ["unable to visit", "unable", "not visited"])
    
    # Add Comments column
    col_comments = None
    if "Comments" in df.columns:
        col_comments = "Comments"
    else:
        for col in df.columns:
            if "comment" in col.lower():
                col_comments = col
                break

    def get_col(colname):
        if colname and colname in df.columns:
            return safe_num(df[colname])
        return pd.Series([0] * len(df), index=df.index, dtype=float)
    
    def get_text_col(colname):
        if colname and colname in df.columns:
            return df[colname].astype(str).fillna("")
        return pd.Series([""] * len(df), index=df.index, dtype=str)

    out = df[["Region", "Province", "District"]].copy()
    out["Total_Sample_Size"] = get_col(col_target)
    out["Total_Received"] = get_col(col_received)
    out["Approved"] = get_col(col_approved)
    out["Pending"] = get_col(col_pending)
    out["Rejected"] = get_col(col_rejected)
    out["Unable_to_Visit"] = get_col(col_unable_visit)
    out["Comments"] = get_text_col(col_comments)

    checked_explicit = get_col(col_checked)
    out["Total_Checked"] = np.where(
        checked_explicit > 0,
        checked_explicit,
        out["Approved"] + out["Pending"] + out["Rejected"]
    )

    if tool == "Total" and out["Total_Sample_Size"].sum() == 0:
        cbe = build_tool_view(df_raw, "CBE")
        pbs = build_tool_view(df_raw, "PBs")
        key = ["Region", "Province", "District"]
        m = cbe.merge(pbs, on=key, how="outer", suffixes=("_CBE", "_PBs")).fillna(0)
        out = m[key].copy()
        out["Total_Sample_Size"] = m["Total_Sample_Size_CBE"] + m["Total_Sample_Size_PBs"]
        out["Total_Received"] = m["Total_Received_CBE"] + m["Total_Received_PBs"]
        out["Approved"] = m["Approved_CBE"] + m["Approved_PBs"]
        out["Pending"] = m["Pending_CBE"] + m["Pending_PBs"]
        out["Rejected"] = m["Rejected_CBE"] + m["Rejected_PBs"]
        out["Total_Checked"] = m["Total_Checked_CBE"] + m["Total_Checked_PBs"]
        out["Unable_to_Visit"] = m.get("Unable_to_Visit_CBE", 0) + m.get("Unable_to_Visit_PBs", 0)
        
        comments_cbe = m.get("Comments_CBE", pd.Series([""] * len(m)))
        comments_pbs = m.get("Comments_PBs", pd.Series([""] * len(m)))
        out["Comments"] = comments_cbe + " | " + comments_pbs

    out["Progress_Percentage"] = np.where(
        out["Total_Sample_Size"] > 0,
        (out["Total_Checked"] / out["Total_Sample_Size"]) * 100.0,
        0
    ).clip(0, 100).round(1)

    out["Progress_Status"] = out["Progress_Percentage"].apply(
        lambda x: "On Track" if x >= 75 else "Behind Schedule" if x >= 50 else "Critical"
    )

    out["_prov_norm"] = out["Province"].map(norm_text)
    out["_dist_norm"] = out["District"].map(norm_text)
    return out

# =========================
# Google Sheet loader
# =========================
@st.cache_data(ttl=600)
def load_google_sheet():
    import gspread
    from google.oauth2.service_account import Credentials

    scopes = ["https://www.googleapis.com/auth/spreadsheets.readonly"]
    
    # Check if secrets are available
    if "gcp_service_account" not in st.secrets:
        st.error("GCP Service Account credentials not found in secrets. Please add them to your Streamlit secrets.")
        return pd.DataFrame()
    
    try:
        creds = Credentials.from_service_account_info(st.secrets["gcp_service_account"], scopes=scopes)
        client = gspread.authorize(creds)

        sh = client.open_by_key(SPREADSHEET_KEY)
        ws = sh.worksheet(WORKSHEET_NAME)
        data = ws.get_all_records()
        df = pd.DataFrame(data)
        df.columns = [str(c).strip() for c in df.columns]
        return df
    except Exception as e:
        st.error(f"Error loading Google Sheet: {e}")
        return pd.DataFrame()

# =========================
# GeoBoundaries Fetchers
# =========================
@st.cache_data(ttl=86400)
def fetch_geoboundaries_geojson(adm_level: str, simplified: bool = True) -> dict:
    adm_level = str(adm_level).upper().strip()
    if adm_level not in ("ADM1", "ADM2"):
        raise ValueError("adm_level must be ADM1 or ADM2")

    api_url = f"https://www.geoboundaries.org/api/current/gbOpen/AFG/{adm_level}/"
    try:
        r = requests.get(api_url, timeout=45)
        r.raise_for_status()
        meta = r.json()
        if isinstance(meta, list) and len(meta) > 0:
            meta = meta[0]

        if simplified:
            geo_url = meta.get("simplifiedGeometryGeoJSON") or meta.get("gjDownloadURL")
        else:
            geo_url = meta.get("gjDownloadURL") or meta.get("simplifiedGeometryGeoJSON")

        if not geo_url:
            return {"type": "FeatureCollection", "features": []}

        g = requests.get(geo_url, timeout=90)
        g.raise_for_status()
        return g.json()
    except Exception as e:
        st.warning(f"Could not fetch GeoBoundaries data: {e}")
        return {"type": "FeatureCollection", "features": []}

def prepare_geojson_for_matching(geojson: dict, name_candidates=("shapeName", "NAME", "NAME_1", "NAME_2", "name")) -> dict:
    for f in geojson.get("features", []):
        props = f.get("properties", {}) or {}
        raw_name = ""
        for k in name_candidates:
            if k in props and props.get(k):
                raw_name = props.get(k)
                break
        props["name_norm"] = norm_text(raw_name)
        f["properties"] = props
    return geojson

# =========================
# Excel Report Functions
# =========================
def create_excel_report(df: pd.DataFrame, level: str = "district", tool_choice: str = "Total", filters: dict = None) -> bytes:
    """
    Create Excel reports at different aggregation levels
    level: 'region', 'province', or 'district'
    """
    output = BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        if level == "region":
            summary = df.groupby("Region").agg({
                "Total_Sample_Size": "sum",
                "Total_Received": "sum",
                "Total_Checked": "sum",
                "Approved": "sum",
                "Pending": "sum",
                "Rejected": "sum",
                "Unable_to_Visit": "sum"
            }).reset_index()
            
            # Calculate percentages
            summary["Progress_Percentage"] = (summary["Total_Checked"] / summary["Total_Sample_Size"] * 100).round(1)
            summary["Approval_Rate"] = (summary["Approved"] / summary["Total_Checked"] * 100).round(1)
            summary["Rejection_Rate"] = (summary["Rejected"] / summary["Total_Checked"] * 100).round(1)
            summary["Received_Rate"] = (summary["Total_Received"] / summary["Total_Sample_Size"] * 100).round(1)
            
            summary.to_excel(writer, sheet_name='Regional_Summary', index=False)
            
        elif level == "province":
            summary = df.groupby(["Region", "Province"]).agg({
                "Total_Sample_Size": "sum",
                "Total_Received": "sum",
                "Total_Checked": "sum",
                "Approved": "sum",
                "Pending": "sum",
                "Rejected": "sum",
                "Unable_to_Visit": "sum",
                "District": "nunique"
            }).reset_index()
            
            summary = summary.rename(columns={"District": "District_Count"})
            summary["Progress_Percentage"] = (summary["Total_Checked"] / summary["Total_Sample_Size"] * 100).round(1)
            summary["Approval_Rate"] = (summary["Approved"] / summary["Total_Checked"] * 100).round(1)
            summary["Rejection_Rate"] = (summary["Rejected"] / summary["Total_Checked"] * 100).round(1)
            summary["Coverage_Rate"] = (summary["Total_Received"] / summary["Total_Sample_Size"] * 100).round(1)
            
            # Add status classification
            summary["Status"] = summary["Progress_Percentage"].apply(
                lambda x: "On Track" if x >= 75 else "Behind Schedule" if x >= 50 else "Critical"
            )
            
            summary.to_excel(writer, sheet_name='Provincial_Summary', index=False)
            
        else:  # district level
            summary = df[["Region", "Province", "District", 
                         "Total_Sample_Size", "Total_Received", "Total_Checked",
                         "Approved", "Pending", "Rejected", "Unable_to_Visit",
                         "Progress_Percentage", "Progress_Status", "Comments"]].copy()
            
            summary["Approval_Rate"] = (summary["Approved"] / summary["Total_Checked"] * 100).round(1)
            summary["Rejection_Rate"] = (summary["Rejected"] / summary["Total_Checked"] * 100).round(1)
            summary["Received_Rate"] = (summary["Total_Received"] / summary["Total_Sample_Size"] * 100).round(1)
            
            summary.to_excel(writer, sheet_name='District_Details', index=False)
        
        # Add metadata sheet
        metadata = pd.DataFrame({
            "Report_Type": [f"{level.title()} Level Report"],
            "Generated_Date": [datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
            "Total_Records": [len(df)],
            "Tool_Used": [tool_choice],
            "Filters_Applied": [str(filters) if filters else "No filters"]
        })
        metadata.to_excel(writer, sheet_name='Metadata', index=False)
    
    output.seek(0)
    return output.getvalue()

# =========================
# Comprehensive Word Report
# =========================
def create_comprehensive_word_report(
    tool_choice: str,
    filters: dict,
    kpis: dict,
    df: pd.DataFrame,
    regional_summary: pd.DataFrame,
    province_summary: pd.DataFrame,
    district_summary: pd.DataFrame,
    unable_to_visit_summary: pd.DataFrame,
    comments_text: str
) -> bytes:
    """
    Create a comprehensive Word report with detailed analysis
    """
    doc = Document()
    
    # Set document properties
    doc.core_properties.author = "Sample Track Analytics System"
    doc.core_properties.title = f"Comprehensive Monitoring Report - {tool_choice}"
    doc.core_properties.subject = "Sample Tracking and Monitoring Analysis"
    doc.core_properties.keywords = "Monitoring, Sample, Tracking, Analytics"
    doc.core_properties.comments = "Comprehensive report generated by Sample Track Analytics Dashboard"
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ========== TITLE PAGE ==========
    # Title
    title_para = doc.add_heading('COMPREHENSIVE SAMPLE TRACKING ANALYSIS REPORT', 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.runs[0].font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Monitoring Tool: {tool_choice}\n")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x47, 0x56, 0x69)
    
    run = subtitle.add_run(f"Date: {datetime.now().strftime('%d %B %Y')}\n")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
    
    doc.add_paragraph()
    
    # Confidential notice
    confidential = doc.add_paragraph()
    confidential.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = confidential.add_run("CONFIDENTIAL - FOR INTERNAL USE ONLY")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)
    run.bold = True
    
    doc.add_page_break()
    
    # ========== TABLE OF CONTENTS ==========
    toc_title = doc.add_heading('TABLE OF CONTENTS', 1)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    toc_items = [
        ("1. EXECUTIVE SUMMARY", 1),
        ("2. METHODOLOGY AND DATA SOURCES", 2),
        ("3. SCOPE AND COVERAGE ANALYSIS", 3),
        ("4. DETAILED PERFORMANCE ANALYSIS", 4),
        ("5. GEOGRAPHICAL DISTRIBUTION", 5),
        ("6. UNABLE TO VISIT ANALYSIS", 6),
        ("7. QUALITY ASSURANCE METRICS", 7),
        ("8. TRENDS AND PATTERNS", 8),
        ("9. CHALLENGES AND OBSERVATIONS", 9),
        ("10. RECOMMENDATIONS", 10),
        ("11. APPENDICES", 11),
        ("Appendix A: Regional Performance Details", 12),
        ("Appendix B: Provincial Performance Details", 13),
        ("Appendix C: District Level Data", 14),
        ("Appendix D: Unable to Visit Details", 15)
    ]
    
    for item, level in toc_items:
        if level == 1:
            para = doc.add_paragraph(item, style='Heading 1')
        elif level == 2:
            para = doc.add_paragraph(f"   {item}", style='Normal')
        else:
            para = doc.add_paragraph(f"      {item}", style='Normal')
    
    doc.add_page_break()
    
    # ========== 1. EXECUTIVE SUMMARY ==========
    doc.add_heading('1. EXECUTIVE SUMMARY', level=1)
    
    exec_summary = doc.add_paragraph()
    exec_summary.add_run("This report provides a comprehensive analysis of the sample tracking and monitoring activities ")
    exec_summary.add_run(f"for the {tool_choice} tool. ").bold = True
    exec_summary.add_run("The analysis covers geographical coverage, performance metrics, quality assurance indicators, ")
    exec_summary.add_run("and operational challenges encountered during the monitoring period.\n\n")
    
    # Key findings
    doc.add_heading('Key Findings', level=2)
    
    findings = [
        f"• Overall Progress: {kpis.get('overall_progress', 0):.1f}% of target samples have been checked",
        f"• Geographical Coverage: {kpis.get('province_count', 0)} provinces and {kpis.get('district_count', 0)} districts covered",
        f"• Quality Rate: {kpis.get('approval_rate', 0):.1f}% approval rate indicates data quality standards",
        f"• Collection Rate: {kpis.get('collection_rate', 0):.1f}% of targeted samples have been received",
        f"• Critical Areas: {len(df[df['Progress_Percentage'] < 50]) if not df.empty else 0} districts are below 50% progress"
    ]
    
    for finding in findings:
        doc.add_paragraph(finding, style='List Bullet')
    
    # ========== 2. METHODOLOGY ==========
    doc.add_heading('2. METHODOLOGY AND DATA SOURCES', level=1)
    
    methodology = doc.add_paragraph()
    methodology.add_run("2.1 Data Collection\n").bold = True
    methodology.add_run("• Primary data source: Google Sheets integrated monitoring tool\n")
    methodology.add_run("• Data extraction: Automated daily synchronization\n")
    methodology.add_run("• Validation: Automated data validation and cleaning procedures\n\n")
    
    methodology.add_run("2.2 Analysis Framework\n").bold = True
    methodology.add_run("• Progress Calculation: (Checked Samples / Target Samples) × 100\n")
    methodology.add_run("• Status Classification:\n")
    methodology.add_run("   - On Track: ≥75% progress\n")
    methodology.add_run("   - Behind Schedule: 50-74% progress\n")
    methodology.add_run("   - Critical: <50% progress\n\n")
    
    methodology.add_run("2.3 Geographical Mapping\n").bold = True
    methodology.add_run("• ADM1 boundaries: Province-level mapping using GeoBoundaries API\n")
    methodology.add_run("• ADM2 boundaries: District-level mapping for detailed analysis\n")
    methodology.add_run("• Normalization: Standardized geographical name matching\n")
    
    # ========== 3. SCOPE AND COVERAGE ==========
    doc.add_heading('3. SCOPE AND COVERAGE ANALYSIS', level=1)
    
    # Coverage table
    coverage_data = [
        ["Metric", "Value", "Interpretation"],
        ["Total Target Samples", f"{kpis.get('total_sample', 0):,.0f}", "Planned sample size across all locations"],
        ["Samples Received", f"{kpis.get('total_received', 0):,.0f}", f"{kpis.get('collection_rate', 0):.1f}% of target"],
        ["Samples Checked", f"{kpis.get('total_checked', 0):,.0f}", f"{kpis.get('overall_progress', 0):.1f}% of target"],
        ["Provinces Covered", f"{kpis.get('province_count', 0)}", "Geographical reach at province level"],
        ["Districts Covered", f"{kpis.get('district_count', 0)}", "Operational presence at district level"],
        ["Approval Rate", f"{kpis.get('approval_rate', 0):.1f}%", "Quality assurance indicator"],
        ["Rejection Rate", f"{kpis.get('rejection_rate', 0):.1f}%", "Quality control measure"]
    ]
    
    table = doc.add_table(rows=len(coverage_data), cols=3)
    table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(coverage_data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_data)
            if i == 0:  # Header row
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    # ========== 4. DETAILED PERFORMANCE ==========
    doc.add_heading('4. DETAILED PERFORMANCE ANALYSIS', level=1)
    
    # Regional performance table
    if regional_summary is not None and not regional_summary.empty:
        doc.add_heading('4.1 Regional Performance Overview', level=2)
        
        reg_data = [["Region", "Target", "Checked", "Progress %", "Status"]]
        for _, row in regional_summary.iterrows():
            progress = row.get("Progress", 0)
            status = "On Track" if progress >= 75 else "Behind Schedule" if progress >= 50 else "Critical"
            reg_data.append([
                row.get("Region", ""),
                f"{row.get('Total_Sample_Size', 0):,.0f}",
                f"{row.get('Total_Checked', 0):,.0f}",
                f"{progress:.1f}%",
                status
            ])
        
        table = doc.add_table(rows=len(reg_data), cols=5)
        table.style = 'Medium Shading 1 Accent 1'
        
        for i, row_data in enumerate(reg_data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(cell_data)
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    
    # Provincial performance
    if province_summary is not None and not province_summary.empty:
        doc.add_heading('4.2 Top Performing Provinces', level=2)
        doc.add_paragraph("The following provinces show the highest progress rates:")
        
        top_provinces = province_summary.sort_values("Progress", ascending=False).head(10)
        prov_data = [["Province", "Progress %", "Target", "Checked", "Approved", "Approval Rate"]]
        
        for _, row in top_provinces.iterrows():
            approval_rate = (row.get("Approved", 0) / row.get("Total_Checked", 1) * 100) if row.get("Total_Checked", 0) > 0 else 0
            prov_data.append([
                row.get("Province", ""),
                f"{row.get('Progress', 0):.1f}%",
                f"{row.get('Total_Sample_Size', 0):,.0f}",
                f"{row.get('Total_Checked', 0):,.0f}",
                f"{row.get('Approved', 0):,.0f}",
                f"{approval_rate:.1f}%"
            ])
        
        table = doc.add_table(rows=len(prov_data), cols=6)
        table.style = 'Light Grid Accent 2'
        
        for i, row_data in enumerate(prov_data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(cell_data)
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    
    # ========== 5. GEOGRAPHICAL DISTRIBUTION ==========
    doc.add_heading('5. GEOGRAPHICAL DISTRIBUTION', level=1)
    
    geo_analysis = doc.add_paragraph()
    geo_analysis.add_run("5.1 Distribution Patterns\n").bold = True
    geo_analysis.add_run(f"The monitoring activities cover {kpis.get('province_count', 0)} provinces across Afghanistan. ")
    geo_analysis.add_run("The geographical distribution shows variations in progress rates, with some regions ")
    geo_analysis.add_run("demonstrating higher efficiency in sample collection and checking processes.\n\n")
    
    geo_analysis.add_run("5.2 Regional Variations\n").bold = True
    geo_analysis.add_run("Analysis indicates significant differences in performance across regions. ")
    geo_analysis.add_run("Factors contributing to these variations include:\n")
    geo_analysis.add_run("• Accessibility and terrain challenges\n")
    geo_analysis.add_run("• Security conditions\n")
    geo_analysis.add_run("• Local capacity and resources\n")
    geo_analysis.add_run("• Logistical constraints\n\n")
    
    # ========== 6. UNABLE TO VISIT ANALYSIS ==========
    doc.add_heading('6. UNABLE TO VISIT ANALYSIS', level=1)
    
    if unable_to_visit_summary is not None and not unable_to_visit_summary.empty:
        doc.add_paragraph(f"A total of {unable_to_visit_summary['Unable_to_Visit'].sum():,.0f} locations were reported as 'Unable to Visit'. ")
        doc.add_paragraph("Primary reasons include:")
        
        reasons = [
            "• Security restrictions and access limitations",
            "• Logistical challenges and transportation issues",
            "• Weather conditions and seasonal factors",
            "• Administrative and permission requirements"
        ]
        
        for reason in reasons:
            doc.add_paragraph(reason)
        
        doc.add_heading('6.1 Detailed Unable to Visit Locations', level=2)
        
        uv_data = [["Province", "District", "Count", "Comments"]]
        for _, row in unable_to_visit_summary.iterrows():
            uv_data.append([
                row.get("Province", ""),
                row.get("District", ""),
                str(int(row.get("Unable_to_Visit", 0))),
                str(row.get("Comments", ""))[:100] + "..." if len(str(row.get("Comments", ""))) > 100 else str(row.get("Comments", ""))
            ])
        
        table = doc.add_table(rows=len(uv_data), cols=4)
        table.style = 'Light List Accent 3'
        
        for i, row_data in enumerate(uv_data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(cell_data)
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    else:
        doc.add_paragraph("No locations were reported as 'Unable to Visit' for the selected filters.")
    
    # ========== 7. QUALITY ASSURANCE ==========
    doc.add_heading('7. QUALITY ASSURANCE METRICS', level=1)
    
    quality = doc.add_paragraph()
    quality.add_run("7.1 Approval and Rejection Rates\n").bold = True
    quality.add_run(f"The overall approval rate of {kpis.get('approval_rate', 0):.1f}% indicates ")
    quality.add_run("acceptable data quality standards. The rejection rate serves as a quality control ")
    quality.add_run("measure to ensure data accuracy and reliability.\n\n")
    
    quality.add_run("7.2 Quality Control Procedures\n").bold = True
    quality.add_run("• Standardized verification protocols\n")
    quality.add_run("• Multi-level review processes\n")
    quality.add_run("• Data validation checks\n")
    quality.add_run("• Consistency verification\n")
    quality.add_run("• Timeliness assessment\n\n")
    
    # ========== 8. TRENDS AND PATTERNS ==========
    doc.add_heading('8. TRENDS AND PATTERNS', level=1)
    
    trends = doc.add_paragraph()
    trends.add_run("8.1 Performance Trends\n").bold = True
    trends.add_run("Analysis reveals several key trends:\n")
    trends.add_run("• Correlation between accessibility and completion rates\n")
    trends.add_run("• Seasonal variations in data collection efficiency\n")
    trends.add_run("• Impact of local capacity on quality metrics\n")
    trends.add_run("• Resource allocation effectiveness\n\n")
    
    trends.add_run("8.2 Pattern Recognition\n").bold = True
    trends.add_run("• Urban areas generally show higher progress rates\n")
    trends.add_run("• Remote districts face greater challenges\n")
    trends.add.run("• Regional coordination impacts overall performance\n")
    
    # ========== 9. CHALLENGES AND OBSERVATIONS ==========
    doc.add_heading('9. CHALLENGES AND OBSERVATIONS', level=1)
    
    if comments_text and comments_text.strip():
        doc.add_paragraph("Key observations from field reports:")
        doc.add_paragraph(comments_text)
    else:
        doc.add_paragraph("No specific observations recorded for the selected filters.")
    
    challenges = doc.add_paragraph("\n\n")
    challenges.add_run("9.1 Common Challenges\n").bold = True
    common_challenges = [
        "• Security constraints limiting access to certain areas",
        "• Logistical challenges in remote and mountainous regions",
        "• Resource limitations affecting monitoring frequency",
        "• Communication barriers in some districts",
        "• Seasonal weather impacts on field operations"
    ]
    
    for challenge in common_challenges:
        doc.add_paragraph(challenge)
    
    # ========== 10. RECOMMENDATIONS ==========
    doc.add_heading('10. RECOMMENDATIONS', level=1)
    
    recommendations = [
        ("10.1 Operational Improvements", [
            "• Increase monitoring frequency in critical areas",
            "• Enhance logistical support for remote districts",
            "• Strengthen local capacity through targeted training",
            "• Implement mobile data collection solutions"
        ]),
        ("10.2 Quality Enhancement", [
            "• Standardize verification protocols across regions",
            "• Implement real-time data validation checks",
            "• Establish quality benchmarks for different regions",
            "• Conduct regular quality assurance audits"
        ]),
        ("10.3 Strategic Planning", [
            "• Develop region-specific action plans",
            "• Allocate resources based on performance indicators",
            "• Establish early warning systems for at-risk areas",
            "• Enhance coordination between regional offices"
        ])
    ]
    
    for title, items in recommendations:
        doc.add_heading(title, level=2)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    # ========== 11. APPENDICES ==========
    doc.add_heading('11. APPENDICES', level=1)
    
    appendices = doc.add_paragraph()
    appendices.add_run("The following appendices provide detailed supporting data for this analysis:\n\n")
    
    appendix_items = [
        ("Appendix A", "Regional Performance Details"),
        ("Appendix B", "Provincial Performance Details"),
        ("Appendix C", "District Level Data"),
        ("Appendix D", "Unable to Visit Details"),
        ("Appendix E", "Methodology Documentation"),
        ("Appendix F", "Quality Assurance Framework")
    ]
    
    for num, title in appendix_items:
        doc.add_paragraph(f"{num}: {title}")
    
    # ========== APPENDICES DETAILS ==========
    doc.add_page_break()
    doc.add_heading('APPENDIX A: REGIONAL PERFORMANCE DETAILS', level=1)
    
    if regional_summary is not None and not regional_summary.empty:
        reg_details = regional_summary.copy()
        reg_details["Progress_Status"] = reg_details["Progress"].apply(
            lambda x: "On Track" if x >= 75 else "Behind Schedule" if x >= 50 else "Critical"
        )
        
        reg_data = [["Region", "Target", "Received", "Checked", "Approved", "Pending", "Rejected", "Progress %", "Status"]]
        for _, row in reg_details.iterrows():
            reg_data.append([
                row.get("Region", ""),
                f"{row.get('Total_Sample_Size', 0):,.0f}",
                f"{row.get('Total_Received', 0):,.0f}",
                f"{row.get('Total_Checked', 0):,.0f}",
                f"{row.get('Approved', 0):,.0f}",
                f"{row.get('Pending', 0):,.0f}",
                f"{row.get('Rejected', 0):,.0f}",
                f"{row.get('Progress', 0):.1f}%",
                row.get("Progress_Status", "")
            ])
        
        table = doc.add_table(rows=len(reg_data), cols=9)
        table.style = 'Light Grid'
        
        for i, row_data in enumerate(reg_data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(cell_data)
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    
    # Add more appendices similarly...
    
    # ========== FINAL PAGE ==========
    doc.add_page_break()
    
    final_page = doc.add_paragraph()
    final_page.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = final_page.add_run("\n\n--- END OF REPORT ---\n\n")
    run.font.size = Pt(12)
    run.bold = True
    
    run = final_page.add_run(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    run.font.size = Pt(10)
    
    run = final_page.add_run("Sample Track Analytics System\n")
    run.font.size = Pt(10)
    run.italic = True
    
    run = final_page.add_run("CONFIDENTIAL - INTERNAL USE ONLY")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)
    
    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output.getvalue()

# =========================
# Enhanced PDF Report
# =========================
def add_page_number(canvas, doc):
    canvas.saveState()
    canvas.setFont('Helvetica', 8)
    canvas.setFillColor(colors.HexColor("#64748b"))
    page_num = f"Page {doc.page}"
    canvas.drawRightString(doc.width + doc.leftMargin, 10*mm, page_num)
    canvas.restoreState()

def make_pdf_report(tool_choice: str, filters: dict, kpis: dict,
                    regional_summary: pd.DataFrame, province_summary: pd.DataFrame,
                    filtered_df: pd.DataFrame, unable_to_visit_summary: pd.DataFrame,
                    comments_text: str) -> bytes:
    buffer = BytesIO()
    
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2*cm,
        rightMargin=2*cm,
        topMargin=3*cm,
        bottomMargin=2*cm
    )

    styles = getSampleStyleSheet()
    
    styles.add(ParagraphStyle(
        name="TitleMain",
        parent=styles["Title"],
        fontSize=22,
        textColor=colors.HexColor("#0f172a"),
        alignment=TA_CENTER,
        spaceAfter=6,
        fontName="Helvetica-Bold"
    ))
    
    styles.add(ParagraphStyle(
        name="Subtitle",
        parent=styles["Normal"],
        fontSize=11,
        textColor=colors.HexColor("#475569"),
        alignment=TA_CENTER,
        spaceAfter=20
    ))
    
    styles.add(ParagraphStyle(
        name="SectionHeader",
        parent=styles["Heading2"],
        fontSize=14,
        textColor=colors.HexColor("#0f172a"),
        spaceBefore=16,
        spaceAfter=10,
        fontName="Helvetica-Bold"
    ))
    
    styles.add(ParagraphStyle(
        name="TableHeader",
        parent=styles["Normal"],
        fontSize=9,
        textColor=colors.white,
        alignment=TA_CENTER,
        fontName="Helvetica-Bold"
    ))
    
    styles.add(ParagraphStyle(
        name="TableCell",
        parent=styles["Normal"],
        fontSize=8,
        textColor=colors.HexColor("#334155"),
        alignment=TA_CENTER
    ))
    
    styles.add(ParagraphStyle(
        name="TableCellLeft",
        parent=styles["Normal"],
        fontSize=8,
        textColor=colors.HexColor("#334155"),
        alignment=TA_LEFT
    ))
    
    styles.add(ParagraphStyle(
        name="Comments",
        parent=styles["Normal"],
        fontSize=9,
        textColor=colors.HexColor("#334155"),
        alignment=TA_LEFT,
        backColor=colors.HexColor("#f8fafc"),
        borderPadding=8,
        spaceBefore=8,
        spaceAfter=8
    ))
    
    story = []
    now_txt = datetime.now().strftime("%Y-%m-%d %H:%M")
    report_date = datetime.now().strftime("%d %B %Y")
    
    # Logo (try to load from theme/assets/logo/ppc.png)
    logo_path = "theme/assets/logo/ppc.png"
    logo_img = None
    if os.path.exists(logo_path):
        try:
            logo_img = Image(logo_path, width=80, height=40)
            logo_img.hAlign = 'LEFT'
        except:
            logo_img = None
    
    # Header with logo and title
    if logo_img:
        header_table = Table([[logo_img, 
                             Paragraph(f"<b>SAMPLE TRACK ANALYTICS REPORT</b><br/>"
                                      f"Monitoring Tool: {tool_choice}<br/>"
                                      f"Date: {report_date}", 
                                      styles["Subtitle"])]], 
                           colWidths=[4*cm, 13*cm])
        header_table.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('ALIGN', (0,0), (0,0), 'LEFT'),
            ('ALIGN', (1,0), (1,0), 'CENTER'),
        ]))
        story.append(header_table)
    else:
        story.append(Paragraph("SAMPLE TRACK ANALYTICS REPORT", styles["TitleMain"]))
        story.append(Paragraph(f"Monitoring Tool: {tool_choice} | Date: {report_date}", styles["Subtitle"]))
    
    story.append(Spacer(1, 0.4*cm))
    
    # Filters Section
    story.append(Paragraph("FILTERS APPLIED", styles["SectionHeader"]))
    f_rows = [
        ["Region", filters.get("region", "All")],
        ["Province", filters.get("province", "All")],
        ["District(s)", filters.get("district", "All")],
        ["Progress Status", ", ".join(filters.get("status", ["All"]))],
    ]
    ft = Table([["Filter", "Value"]] + f_rows, colWidths=[5*cm, 10*cm])
    ft.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#0f172a")),
        ("TEXTCOLOR", (0,0), (-1,0), colors.white),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE", (0,0), (-1,-1), 9),
        ("GRID", (0,0), (-1,-1), 0.5, colors.HexColor("#cbd5e1")),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("PADDING", (0,0), (-1,-1), 6),
    ]))
    story.append(ft)
    story.append(Spacer(1, 0.8*cm))
    
    # Executive Summary
    story.append(Paragraph("EXECUTIVE SUMMARY", styles["SectionHeader"]))
    es_data = [
        ["Metric", "Value", "Details"],
        ["Total Target", f"{kpis.get('total_sample', 0):,.0f}", "Planned sample size"],
        ["Total Received", f"{kpis.get('total_received', 0):,.0f}", "Samples collected"],
        ["Total Checked", f"{kpis.get('total_checked', 0):,.0f}", "Samples reviewed"],
        ["Approved", f"{kpis.get('total_approved', 0):,.0f}", f"Approval Rate: {kpis.get('approval_rate', 0):.1f}%"],
        ["Rejected", f"{kpis.get('total_rejected', 0):,.0f}", f"Rejection Rate: {kpis.get('rejection_rate', 0):.1f}%"],
        ["Pending", f"{kpis.get('total_pending', 0):,.0f}", "Awaiting review"],
        ["Overall Progress", f"{kpis.get('overall_progress', 0):.1f}%", "Checked vs Target"],
        ["Coverage", f"{kpis.get('province_count', 0)} Prov / {kpis.get('district_count', 0)} Dist", "Geographical coverage"],
    ]
    es_table = Table(es_data, colWidths=[5*cm, 3.5*cm, 6.5*cm])
    es_table.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#1e40af")),
        ("TEXTCOLOR", (0,0), (-1,0), colors.white),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE", (0,0), (-1,-1), 9),
        ("GRID", (0,0), (-1,-1), 0.5, colors.HexColor("#cbd5e1")),
        ("BACKGROUND", (0,1), (-1,-1), colors.HexColor("#ffffff")),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("PADDING", (0,0), (-1,-1), 6),
    ]))
    story.append(es_table)
    
    story.append(Spacer(1, 0.8*cm))
    
    # Unable to Visit Section
    if unable_to_visit_summary is not None and not unable_to_visit_summary.empty:
        story.append(Paragraph(f"UNABLE TO VISIT - {tool_choice}", styles["SectionHeader"]))
        
        uv_data = [["Province", "District", "Unable to Visit Count", "Comments"]]
        for _, row in unable_to_visit_summary.iterrows():
            comments = str(row.get("Comments", "")).strip()
            if len(comments) > 100:
                comments = comments[:97] + "..."
            uv_data.append([
                row.get("Province", ""),
                row.get("District", ""),
                str(int(row.get("Unable_to_Visit", 0))),
                comments
            ])
        
        uv_table = Table(uv_data, colWidths=[4*cm, 4*cm, 3*cm, 8*cm])
        uv_table.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#dc2626")),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE", (0,0), (-1,-1), 8),
            ("GRID", (0,0), (-1,-1), 0.5, colors.HexColor("#cbd5e1")),
            ("BACKGROUND", (0,1), (-1,-1), colors.HexColor("#fef2f2")),
            ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
            ("PADDING", (0,0), (-1,-1), 5),
            ("ALIGN", (2,1), (2,-1), "CENTER"),
        ]))
        story.append(uv_table)
        story.append(Spacer(1, 0.8*cm))
    
    # Comments Section
    if comments_text and comments_text.strip():
        story.append(Paragraph("KEY COMMENTS & OBSERVATIONS", styles["SectionHeader"]))
        story.append(Paragraph(comments_text, styles["Comments"]))
        story.append(Spacer(1, 0.8*cm))
    
    # Regional Performance
    story.append(Paragraph("REGIONAL PERFORMANCE", styles["SectionHeader"]))
    if regional_summary is not None and not regional_summary.empty:
        reg_data = [["Region", "Target", "Checked", "Progress %", "Status"]]
        for _, row in regional_summary.iterrows():
            progress = row.get("Progress", 0)
            status = "On Track" if progress >= 75 else "Behind Schedule" if progress >= 50 else "Critical"
            reg_data.append([
                row.get("Region", ""),
                f"{row.get('Total_Sample_Size', 0):,.0f}",
                f"{row.get('Total_Checked', 0):,.0f}",
                f"{progress:.1f}%",
                status
            ])
        
        reg_table = Table(reg_data, colWidths=[4*cm, 3*cm, 3*cm, 3*cm, 4*cm])
        reg_table.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#0f766e")),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE", (0,0), (-1,-1), 9),
            ("GRID", (0,0), (-1,-1), 0.5, colors.HexColor("#cbd5e1")),
            ("BACKGROUND", (0,1), (-1,-1), colors.HexColor("#ffffff")),
            ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
            ("PADDING", (0,0), (-1,-1), 6),
        ]))
        story.append(reg_table)
    else:
        story.append(Paragraph("No regional data available for selected filters.", styles["TableCellLeft"]))
    
    story.append(Spacer(1, 0.8*cm))
    
    # Province Performance (Top 10)
    story.append(Paragraph("TOP 10 PROVINCES BY PROGRESS", styles["SectionHeader"]))
    if province_summary is not None and not province_summary.empty:
        top_provinces = province_summary.sort_values("Progress", ascending=False).head(10)
        prov_data = [["Province", "Progress %", "Target", "Checked", "Approved", "Rejected"]]
        for _, row in top_provinces.iterrows():
            prov_data.append([
                row.get("Province", ""),
                f"{row.get('Progress', 0):.1f}%",
                f"{row.get('Total_Sample_Size', 0):,.0f}",
                f"{row.get('Total_Checked', 0):,.0f}",
                f"{row.get('Approved', 0):,.0f}",
                f"{row.get('Rejected', 0):,.0f}"
            ])
        
        prov_table = Table(prov_data, colWidths=[4*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm])
        prov_table.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#1d4ed8")),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE", (0,0), (-1,-1), 8),
            ("GRID", (0,0), (-1,-1), 0.5, colors.HexColor("#cbd5e1")),
            ("BACKGROUND", (0,1), (-1,-1), colors.HexColor("#ffffff")),
            ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
            ("PADDING", (0,0), (-1,-1), 5),
        ]))
        story.append(prov_table)
    else:
        story.append(Paragraph("No province summary available.", styles["TableCellLeft"]))
    
    story.append(Spacer(1, 1*cm))
    
    # Footer
    story.append(Paragraph(f"Report Generated: {now_txt}", 
                          ParagraphStyle(name="Footer", fontSize=8, textColor=colors.HexColor("#64748b"))))
    story.append(Paragraph("Confidential - For Internal Use Only", 
                          ParagraphStyle(name="Footer", fontSize=8, textColor=colors.HexColor("#64748b"))))
    
    # Build document with page numbers
    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
    return buffer.getvalue()

# =========================
# Load data
# =========================
try:
    df_sheet = load_google_sheet()
    if df_sheet.empty:
        st.error("Google Sheet is empty.")
        st.stop()
except Exception as e:
    st.error(f"Failed to load data from Google Sheet: {e}")
    st.stop()

# =========================
# Sidebar
# =========================
st.sidebar.markdown("### Filters & Controls")

tool_choice = st.sidebar.selectbox(
    "Select Monitoring Tool",
    ["Total", "CBE", "PBs"],
    help="The app reads columns by prefix: CBE-, PBs-, or Total-. If Total- does not exist, Total is calculated as CBE + PBs."
)

try:
    df = build_tool_view(df_sheet, tool_choice)
except Exception as e:
    st.error(f"Error processing data: {e}")
    st.stop()

selected_region = st.sidebar.selectbox("Select Region", ["All"] + sorted(df["Region"].dropna().unique().tolist()))

if selected_region != "All":
    province_options = ["All"] + sorted(df[df["Region"] == selected_region]["Province"].dropna().unique().tolist())
else:
    province_options = ["All"] + sorted(df["Province"].dropna().unique().tolist())
selected_province = st.sidebar.selectbox("Select Province", province_options)

if selected_province != "All":
    district_options = sorted(df[df["Province"] == selected_province]["District"].dropna().unique().tolist())
elif selected_region != "All":
    district_options = sorted(df[df["Region"] == selected_region]["District"].dropna().unique().tolist())
else:
    district_options = sorted(df["District"].dropna().unique().tolist())

selected_districts = st.sidebar.multiselect(
    "Select District(s)",
    options=district_options,
    default=[],
    help="Leave empty to include all districts."
)

progress_status = st.sidebar.multiselect(
    "Progress Status",
    options=["All", "On Track", "Behind Schedule", "Critical"],
    default=["All"]
)

# Apply filters
filtered_df = df.copy()

if selected_region != "All":
    filtered_df = filtered_df[filtered_df["Region"] == selected_region]

if selected_province != "All":
    filtered_df = filtered_df[filtered_df["Province"] == selected_province]

if selected_districts:
    filtered_df = filtered_df[filtered_df["District"].isin(selected_districts)]

if "All" not in progress_status:
    filtered_df = filtered_df[filtered_df["Progress_Status"].isin(progress_status)]

# =========================
# KPIs
# =========================
st.markdown('<div class="section-title">Performance Overview</div>', unsafe_allow_html=True)

total_sample = float(filtered_df["Total_Sample_Size"].sum())
total_received = float(filtered_df["Total_Received"].sum())
total_approved = float(filtered_df["Approved"].sum())
total_pending = float(filtered_df["Pending"].sum())
total_rejected = float(filtered_df["Rejected"].sum())
total_checked = float(filtered_df["Total_Checked"].sum())
total_unable_visit = float(filtered_df["Unable_to_Visit"].sum())

overall_progress = (total_checked / total_sample * 100.0) if total_sample > 0 else 0.0
approval_rate = (total_approved / total_checked * 100.0) if total_checked > 0 else 0.0
rejection_rate = (total_rejected / total_checked * 100.0) if total_checked > 0 else 0.0

c1, c2, c3, c4 = st.columns(4)
with c1:
    st.markdown(f"""
    <div class="kpi-card">
        <div class="kpi-label">Total Target</div>
        <div class="kpi-value">{total_sample:,.0f}</div>
        <div class="kpi-sub">Provinces: {filtered_df["Province"].nunique():,} | Districts: {filtered_df["District"].nunique():,}</div>
    </div>
    """, unsafe_allow_html=True)

with c2:
    st.markdown(f"""
    <div class="kpi-card">
        <div class="kpi-label">Overall Progress</div>
        <div class="kpi-value">{overall_progress:.1f}%</div>
        <div class="kpi-sub">Checked: {total_checked:,.0f} | Remaining: {max(total_sample-total_checked, 0):,.0f}</div>
    </div>
    """, unsafe_allow_html=True)

with c3:
    st.markdown(f"""
    <div class="kpi-card">
        <div class="kpi-label">Approvals</div>
        <div class="kpi-value">{total_approved:,.0f}</div>
        <div class="kpi-sub">Approval Rate: {approval_rate:.1f}%</div>
    </div>
    """, unsafe_allow_html=True)

with c4:
    st.markdown(f"""
    <div class="kpi-card">
        <div class="kpi-label">Rejections / Unable to Visit</div>
        <div class="kpi-value">{total_rejected:,.0f} / {total_unable_visit:,.0f}</div>
        <div class="kpi-sub">Rejection Rate: {rejection_rate:.1f}% | Received: {total_received:,.0f}</div>
    </div>
    """, unsafe_allow_html=True)

# =========================
# Maps (ADM1 + ADM2)
# =========================
st.markdown('<div class="section-title">Geographic Analysis</div>', unsafe_allow_html=True)

with st.spinner("Loading boundary layers (ADM1 and ADM2)..."):
    try:
        adm1_geojson = fetch_geoboundaries_geojson("ADM1", simplified=True)
        adm2_geojson = fetch_geoboundaries_geojson("ADM2", simplified=True)

        adm1_geojson = prepare_geojson_for_matching(adm1_geojson, name_candidates=("shapeName", "NAME_1", "NAME", "name"))
        adm2_geojson = prepare_geojson_for_matching(adm2_geojson, name_candidates=("shapeName", "NAME_2", "NAME", "name"))
    except Exception as e:
        st.error(f"Map layers could not be loaded. {e}")
        adm1_geojson = {"type": "FeatureCollection", "features": []}
        adm2_geojson = {"type": "FeatureCollection", "features": []}

# ADM1: Afghanistan Province Map
st.markdown('<div class="section-title">Afghanistan (ADM1 - Provinces)</div>', unsafe_allow_html=True)

if adm1_geojson.get("features"):
    prov_summary = filtered_df.groupby(["Province", "_prov_norm"], as_index=False).agg({
        "Total_Sample_Size":"sum",
        "Total_Checked":"sum",
        "Approved":"sum",
        "Rejected":"sum",
        "Pending":"sum"
    })

    prov_summary["Progress_Percentage"] = np.where(
        prov_summary["Total_Sample_Size"] > 0,
        (prov_summary["Total_Checked"] / prov_summary["Total_Sample_Size"]) * 100.0,
        0
    ).round(1)

    fig_afg = px.choropleth(
        prov_summary,
        geojson=adm1_geojson,
        locations="_prov_norm",
        featureidkey="properties.name_norm",
        color="Progress_Percentage",
        hover_name="Province",
        hover_data={
            "Progress_Percentage": ":.1f",
            "Total_Sample_Size": ":,.0f",
            "Total_Checked": ":,.0f",
            "Approved": ":,.0f",
            "Rejected": ":,.0f",
            "Pending": ":,.0f",
            "_prov_norm": False
        },
        color_continuous_scale="RdYlGn",
        range_color=[0, 100],
        title=""
    )

    if selected_province == "All":
        fig_afg.update_geos(visible=False, fitbounds=None)
    else:
        fig_afg.update_geos(visible=False, fitbounds="locations")

    fig_afg.update_layout(height=520, margin=dict(l=0, r=0, t=10, b=0))
    st.plotly_chart(fig_afg, use_container_width=True)
else:
    st.warning("ADM1 map could not be loaded.")

# ADM2: District Map
st.markdown('<div class="section-title">Districts (ADM2)</div>', unsafe_allow_html=True)

if not adm2_geojson.get("features"):
    st.warning("ADM2 map could not be loaded.")
else:
    if selected_province != "All":
        df_scope = filtered_df[filtered_df["Province"] == selected_province].copy()
        if df_scope.empty:
            st.info("No district data for selected filters.")
        else:
            dist_summary = df_scope.groupby(["District", "_dist_norm"], as_index=False).agg({
                "Total_Sample_Size":"sum",
                "Total_Checked":"sum",
                "Approved":"sum",
                "Rejected":"sum",
                "Pending":"sum"
            })

            dist_summary["Progress_Percentage"] = np.where(
                dist_summary["Total_Sample_Size"] > 0,
                (dist_summary["Total_Checked"] / dist_summary["Total_Sample_Size"]) * 100.0,
                0
            ).round(1)

            if selected_districts:
                dist_summary = dist_summary[dist_summary["District"].isin(selected_districts)]

            if dist_summary.empty:
                st.info("No matching districts found.")
            else:
                fig_dist = px.choropleth(
                    dist_summary,
                    geojson=adm2_geojson,
                    locations="_dist_norm",
                    featureidkey="properties.name_norm",
                    color="Progress_Percentage",
                    hover_name="District",
                    hover_data={
                        "Progress_Percentage": ":.1f",
                        "Total_Sample_Size": ":,.0f",
                        "Total_Checked": ":,.0f",
                        "Approved": ":,.0f",
                        "Rejected": ":,.0f",
                        "Pending": ":,.0f",
                        "_dist_norm": False
                    },
                    color_continuous_scale="RdYlGn",
                    range_color=[0, 100],
                    title=""
                )
                fig_dist.update_geos(visible=False, fitbounds="locations")
                fig_dist.update_layout(height=520, margin=dict(l=0, r=0, t=10, b=0))
                st.plotly_chart(fig_dist, use_container_width=True)
    else:
        st.markdown(
            '<div class="hint">For faster performance, select a province.</div>',
            unsafe_allow_html=True
        )

        df_scope = filtered_df.copy()
        if df_scope.empty:
            st.info("No data for selected filters.")
        else:
            dist_summary = df_scope.groupby(["District", "_dist_norm"], as_index=False).agg({
                "Total_Sample_Size":"sum",
                "Total_Checked":"sum",
                "Approved":"sum",
                "Rejected":"sum",
                "Pending":"sum"
            })

            dist_summary["Progress_Percentage"] = np.where(
                dist_summary["Total_Sample_Size"] > 0,
                (dist_summary["Total_Checked"] / dist_summary["Total_Sample_Size"]) * 100.0,
                0
            ).round(1)

            if selected_districts:
                dist_summary = dist_summary[dist_summary["District"].isin(selected_districts)]

            if dist_summary.empty:
                st.info("No matching districts found.")
            else:
                fig_dist_all = px.choropleth(
                    dist_summary,
                    geojson=adm2_geojson,
                    locations="_dist_norm",
                    featureidkey="properties.name_norm",
                    color="Progress_Percentage",
                    hover_name="District",
                    hover_data={
                        "Progress_Percentage": ":.1f",
                        "Total_Sample_Size": ":,.0f",
                        "Total_Checked": ":,.0f",
                        "Approved": ":,.0f",
                        "Rejected": ":,.0f",
                        "Pending": ":,.0f",
                        "_dist_norm": False
                    },
                    color_continuous_scale="RdYlGn",
                    range_color=[0, 100],
                    title=""
                )

                if selected_districts:
                    fig_dist_all.update_geos(visible=False, fitbounds="locations")
                else:
                    fig_dist_all.update_geos(visible=False, fitbounds=None)

                fig_dist_all.update_layout(height=560, margin=dict(l=0, r=0, t=10, b=0))
                st.plotly_chart(fig_dist_all, use_container_width=True)

# =========================
# Charts
# =========================
st.markdown('<div class="section-title">Charts</div>', unsafe_allow_html=True)

t1, t2, t3 = st.tabs(["Status Breakdown", "Region Comparison", "Target vs Checked"])

with t1:
    status_data = pd.DataFrame({
        "Category": ["Approved", "Pending", "Rejected", "Unable to Visit", "Not Checked"],
        "Count": [total_approved, total_pending, total_rejected, total_unable_visit, max(total_sample-total_checked, 0)]
    })
    fig_bar = px.bar(status_data, x="Category", y="Count", text="Count")
    fig_bar.update_traces(texttemplate="%{text:,}", textposition="outside")
    fig_bar.update_layout(height=420, xaxis_title="", yaxis_title="Count", showlegend=False)
    st.plotly_chart(fig_bar, use_container_width=True)

with t2:
    regional_summary = filtered_df.groupby("Region", as_index=False).agg({
        "Total_Sample_Size":"sum",
        "Total_Checked":"sum",
        "Approved":"sum",
        "Rejected":"sum",
        "Pending":"sum"
    })
    regional_summary["Progress"] = np.where(
        regional_summary["Total_Sample_Size"] > 0,
        (regional_summary["Total_Checked"] / regional_summary["Total_Sample_Size"]) * 100.0,
        0
    ).round(1)

    fig_region = px.bar(regional_summary, x="Region", y="Progress", text="Progress")
    fig_region.update_traces(texttemplate="%{text:.1f}%", textposition="outside")
    fig_region.update_layout(height=420, xaxis_title="", yaxis_title="Progress (%)")
    st.plotly_chart(fig_region, use_container_width=True)

with t3:
    if not filtered_df.empty:
        fig_scatter = px.scatter(
            filtered_df,
            x="Total_Sample_Size",
            y="Total_Checked",
            size="Total_Sample_Size",
            color="Progress_Percentage",
            hover_name="District",
            hover_data=["Province", "Approved", "Rejected", "Pending", "Unable_to_Visit"],
            size_max=28
        )
        max_val = max(filtered_df["Total_Sample_Size"].max(), filtered_df["Total_Checked"].max())
        fig_scatter.add_trace(go.Scatter(
            x=[0, max_val],
            y=[0, max_val],
            mode="lines",
            name="Target = Checked",
            line=dict(dash="dash")
        ))
        fig_scatter.update_layout(height=520, xaxis_title="Target", yaxis_title="Checked")
        st.plotly_chart(fig_scatter, use_container_width=True)

# =========================
# Tables
# =========================
st.markdown('<div class="section-title">Tables</div>', unsafe_allow_html=True)

left, right = st.columns(2)

with left:
    st.subheader("Province Summary")
    province_summary = filtered_df.groupby("Province", as_index=False).agg({
        "Total_Sample_Size":"sum",
        "Total_Received":"sum",
        "Total_Checked":"sum",
        "Approved":"sum",
        "Rejected":"sum",
        "Pending":"sum",
        "Unable_to_Visit":"sum",
        "Progress_Percentage":"mean"
    })
    province_summary["Progress"] = province_summary["Progress_Percentage"].round(1)
    province_summary["Approval_Rate"] = np.where(
        province_summary["Total_Checked"] > 0,
        (province_summary["Approved"] / province_summary["Total_Checked"]) * 100.0,
        0
    ).round(1)

    st.dataframe(
        province_summary.sort_values("Progress", ascending=False),
        use_container_width=True,
        height=420
    )

with right:
    st.subheader("District Performance")
    district_summary = filtered_df.groupby(["Province", "District"], as_index=False).agg({
        "Total_Sample_Size":"sum",
        "Total_Checked":"sum",
        "Approved":"sum",
        "Rejected":"sum",
        "Pending":"sum",
        "Unable_to_Visit":"sum",
        "Progress_Percentage":"mean"
    }).round(1)

    st.dataframe(
        district_summary.sort_values("Progress_Percentage", ascending=False).head(50),
        use_container_width=True,
        height=420
    )

# =========================
# Prepare data for reports
# =========================
# Prepare Unable to Visit summary
unable_to_visit_summary = None
if "Unable_to_Visit" in filtered_df.columns:
    unable_df = filtered_df[filtered_df["Unable_to_Visit"] > 0]
    if not unable_df.empty:
        unable_to_visit_summary = unable_df[["Province", "District", "Unable_to_Visit", "Comments"]].copy()
        unable_to_visit_summary = unable_to_visit_summary.sort_values("Unable_to_Visit", ascending=False)

# Prepare comments text
comments_text = ""
if "Comments" in filtered_df.columns:
    comments_list = filtered_df["Comments"].dropna().unique()
    meaningful_comments = [str(c).strip() for c in comments_list if str(c).strip() and str(c).strip().lower() not in ["", "nan", "none", "n/a"]]
    if meaningful_comments:
        comments_text = " | ".join(meaningful_comments)
        if len(comments_text) > 2000:
            comments_text = comments_text[:1997] + "..."

# Prepare regional summary for reports
regional_summary_report = filtered_df.groupby("Region", as_index=False).agg({
    "Total_Sample_Size":"sum",
    "Total_Checked":"sum",
    "Approved":"sum",
    "Rejected":"sum",
    "Pending":"sum"
})
regional_summary_report["Progress"] = np.where(
    regional_summary_report["Total_Sample_Size"] > 0,
    (regional_summary_report["Total_Checked"] / regional_summary_report["Total_Sample_Size"]) * 100.0,
    0
).round(1)

# Prepare province summary for reports
province_summary_report = filtered_df.groupby("Province", as_index=False).agg({
    "Total_Sample_Size":"sum",
    "Total_Checked":"sum",
    "Approved":"sum",
    "Rejected":"sum",
    "Pending":"sum"
})
province_summary_report["Progress"] = np.where(
    province_summary_report["Total_Sample_Size"] > 0,
    (province_summary_report["Total_Checked"] / province_summary_report["Total_Sample_Size"]) * 100.0,
    0
).round(1)

# Prepare district summary for reports
district_summary_report = filtered_df.groupby(["Province", "District"], as_index=False).agg({
    "Total_Sample_Size":"sum",
    "Total_Received":"sum",
    "Total_Checked":"sum",
    "Approved":"sum",
    "Pending":"sum",
    "Rejected":"sum",
    "Unable_to_Visit":"sum",
    "Progress_Percentage":"mean",
    "Comments": lambda x: " | ".join([str(i) for i in x if str(i).strip()])
}).round(1)

# Prepare filters dictionary
filters_for_report = {
    "region": selected_region,
    "province": selected_province,
    "district": ", ".join(selected_districts) if selected_districts else "All",
    "status": progress_status
}

# Prepare KPIs dictionary
kpis_for_report = {
    "total_sample": total_sample,
    "total_received": total_received,
    "total_checked": total_checked,
    "total_approved": total_approved,
    "total_rejected": total_rejected,
    "total_pending": total_pending,
    "overall_progress": overall_progress,
    "approval_rate": approval_rate,
    "rejection_rate": rejection_rate,
    "province_count": filtered_df["Province"].nunique(),
    "district_count": filtered_df["District"].nunique()
}

# =========================
# Enhanced Report Export Section
# =========================
st.markdown('<div class="section-title">Comprehensive Report Export</div>', unsafe_allow_html=True)

# Create tabs for different export options
export_tab1, export_tab2, export_tab3 = st.tabs(["📊 Excel Downloads", "📝 Word Report", "📄 PDF Report"])

with export_tab1:
    st.subheader("Download Excel Reports")
    st.markdown("""
    <div class="hint">
    Download detailed Excel reports at different geographical levels. Each report includes 
    comprehensive metrics, calculations, and analysis-ready data.
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("📥 Regional Summary", use_container_width=True, key="excel_region"):
            with st.spinner("Generating Regional Excel Report..."):
                try:
                    excel_bytes = create_excel_report(filtered_df, "region", tool_choice, filters_for_report)
                    filename = f"Regional_Summary_{tool_choice}_{datetime.now().strftime('%Y%m%d')}.xlsx"
                    
                    st.download_button(
                        label="Download Regional Excel",
                        data=excel_bytes,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True,
                        key="download_region"
                    )
                except Exception as e:
                    st.error(f"Error generating regional report: {e}")
    
    with col2:
        if st.button("📥 Provincial Summary", use_container_width=True, key="excel_province"):
            with st.spinner("Generating Provincial Excel Report..."):
                try:
                    excel_bytes = create_excel_report(filtered_df, "province", tool_choice, filters_for_report)
                    filename = f"Provincial_Summary_{tool_choice}_{datetime.now().strftime('%Y%m%d')}.xlsx"
                    
                    st.download_button(
                        label="Download Provincial Excel",
                        data=excel_bytes,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True,
                        key="download_province"
                    )
                except Exception as e:
                    st.error(f"Error generating provincial report: {e}")
    
    with col3:
        if st.button("📥 District Details", use_container_width=True, key="excel_district"):
            with st.spinner("Generating District Excel Report..."):
                try:
                    excel_bytes = create_excel_report(filtered_df, "district", tool_choice, filters_for_report)
                    filename = f"District_Details_{tool_choice}_{datetime.now().strftime('%Y%m%d')}.xlsx"
                    
                    st.download_button(
                        label="Download District Excel",
                        data=excel_bytes,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True,
                        key="download_district"
                    )
                except Exception as e:
                    st.error(f"Error generating district report: {e}")
    
    # Additional Excel options
    st.markdown("---")
    st.subheader("Custom Excel Downloads")
    
    custom_col1, custom_col2 = st.columns(2)
    
    with custom_col1:
        if st.button("📊 Performance Metrics Only", use_container_width=True, key="excel_performance"):
            with st.spinner("Generating Performance Metrics..."):
                try:
                    # Create simplified performance metrics
                    perf_df = filtered_df[["Region", "Province", "District", 
                                          "Progress_Percentage", "Progress_Status",
                                          "Total_Sample_Size", "Total_Checked"]].copy()
                    output = BytesIO()
                    with pd.ExcelWriter(output, engine='openpyxl') as writer:
                        perf_df.to_excel(writer, sheet_name='Performance_Metrics', index=False)
                    output.seek(0)
                    
                    st.download_button(
                        label="Download Performance Data",
                        data=output.getvalue(),
                        file_name=f"Performance_Metrics_{tool_choice}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True,
                        key="download_performance"
                    )
                except Exception as e:
                    st.error(f"Error generating performance report: {e}")
    
    with custom_col2:
        if st.button("⚠️ Critical Areas Report", use_container_width=True, key="excel_critical"):
            with st.spinner("Generating Critical Areas Report..."):
                try:
                    # Filter critical areas
                    critical_df = filtered_df[filtered_df["Progress_Percentage"] < 50].copy()
                    if not critical_df.empty:
                        output = BytesIO()
                        with pd.ExcelWriter(output, engine='openpyxl') as writer:
                            critical_df.to_excel(writer, sheet_name='Critical_Areas', index=False)
                        output.seek(0)
                        
                        st.download_button(
                            label="Download Critical Areas",
                            data=output.getvalue(),
                            file_name=f"Critical_Areas_{tool_choice}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True,
                            key="download_critical"
                        )
                    else:
                        st.info("No critical areas (below 50% progress) found with current filters.")
                except Exception as e:
                    st.error(f"Error generating critical areas report: {e}")

with export_tab2:
    st.subheader("Generate Comprehensive Word Report")
    st.markdown("""
    <div class="hint">
    This comprehensive Word report includes detailed analysis, methodology, findings, 
    recommendations, and appendices with supporting data. The report is professionally 
    formatted and ready for presentation.
    </div>
    """, unsafe_allow_html=True)
    
    if st.button("📝 Generate Comprehensive Word Report", type="primary", use_container_width=True, key="word_report"):
        with st.spinner("Generating comprehensive Word report..."):
            try:
                # Calculate additional KPIs for Word report
                kpis_for_word = kpis_for_report.copy()
                kpis_for_word["collection_rate"] = (kpis_for_word["total_received"] / kpis_for_word["total_sample"] * 100) if kpis_for_word["total_sample"] > 0 else 0
                
                # Generate Word report
                word_bytes = create_comprehensive_word_report(
                    tool_choice=tool_choice,
                    filters=filters_for_report,
                    kpis=kpis_for_word,
                    df=filtered_df,
                    regional_summary=regional_summary_report,
                    province_summary=province_summary_report,
                    district_summary=district_summary_report,
                    unable_to_visit_summary=unable_to_visit_summary,
                    comments_text=comments_text
                )
                
                filename = f"Comprehensive_Report_{tool_choice}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                
                st.download_button(
                    label="📥 Download Word Report",
                    data=word_bytes,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="download_word"
                )
                
                st.success("✅ Comprehensive Word report generated successfully!")
                st.info("The report includes: Executive Summary, Methodology, Detailed Analysis, Recommendations, and Appendices.")
                
            except Exception as e:
                st.error(f"Error generating Word report: {e}")

with export_tab3:
    st.subheader("Generate PDF Report")
    st.markdown("""
    <div class="hint">
    Generate a concise PDF report with key metrics, charts, and summary information.
    This is ideal for quick sharing and presentations.
    </div>
    """, unsafe_allow_html=True)
    
    if st.button("📄 Generate PDF Report", type="primary", use_container_width=True, key="pdf_report"):
        with st.spinner("Generating PDF report..."):
            try:
                pdf_bytes = make_pdf_report(
                    tool_choice=tool_choice,
                    filters=filters_for_report,
                    kpis=kpis_for_report,
                    regional_summary=regional_summary_report,
                    province_summary=province_summary_report,
                    filtered_df=filtered_df,
                    unable_to_visit_summary=unable_to_visit_summary,
                    comments_text=comments_text
                )

                filename = f"SampleTrack_Report_{tool_choice}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
                
                st.download_button(
                    label="📥 Download PDF Report",
                    data=pdf_bytes,
                    file_name=filename,
                    mime="application/pdf",
                    use_container_width=True,
                    key="download_pdf"
                )
                
                st.success("✅ PDF report generated successfully!")
                
            except Exception as e:
                st.error(f"Error generating PDF: {e}")

# =========================
# Footer
# =========================
st.markdown("---")
st.caption(
    f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | "
    f"Tool: {tool_choice} | Records: {len(filtered_df):,} | "
    f"Unable to Visit: {total_unable_visit:,.0f}"
)
